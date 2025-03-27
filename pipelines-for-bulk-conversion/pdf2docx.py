"""
this python script is an end-to-end pdf-to-docx converter.
instructions to use:
- using requirements.txt, install all libraries
- in the directory where this script is present, create a directory - `pdf-docx`
- in `pdf-docx`, create a directory - `to-convert`
- transfer all the pdfs that need to be converted to this directory (i.e. `pdf-docx/to-convert`)
- run this script! your outputs will be organised in `pdf-docx/working/{pdf_name}`
"""


### importing libraries
import os, re
from docx import Document as docx_document
from spire.pdf.common import *
from spire.pdf import *
from spire.pdf import FileFormat as FF
from spire.doc.common import *
from spire.doc import *

### step 1 - splitting PDFs
def split_pdf(pdf_path, pdf_name):
    # creating a directory to store the split PDFs
    os.makedirs(f"pdf-docx/working/{pdf_name}/split_pdfs", exist_ok = True)
    # creating a PdfDocument object
    target_doc = PdfDocument()
    # loading target PDF file
    target_doc.LoadFromFile(pdf_path)

    # creating PdfDocument objects that will consist of the split target PDF
    num_required_splits = target_doc.Pages.Count // 10 + 1
    docs = [PdfDocument() for _ in range(num_required_splits)]
    
    doc_index = 0
    remainingPages = target_doc.Pages.Count
    current_page = 0
    while remainingPages >= 1:
        if remainingPages >= 10:
            docs[doc_index].InsertPageRange(target_doc, current_page, current_page + 9)
            current_page += 10
            remainingPages -= 10
            doc_index += 1
            continue
        elif remainingPages >= 1:
            docs[doc_index].InsertPageRange(target_doc, current_page, current_page + remainingPages - 1)
            doc_index += 1
            remainingPages = 0
            current_page += remainingPages
            break
    
    for i, doc in enumerate(docs):
        doc.SaveToFile(f"pdf-docx/working/{pdf_name}/split_pdfs/Split-{i + 1}.pdf")
        doc.Close() # closing the PdfDocument object

    target_doc.Close()
    return f"pdf-docx/working/{pdf_name}/split_pdfs" # return the path to the directory containing the split PDFs

### step 2 - converting split PDFs to DOCXs
def convert_pdf(split_pdf_dir, pdf_name):
    os.makedirs(f"pdf-docx/working/{pdf_name}/split_docxs", exist_ok = True)
    for split_pdf_path in os.listdir(split_pdf_dir):
        pdf = PdfDocument()
        pdf.LoadFromFile(os.path.join(split_pdf_dir, split_pdf_path))
        
        # converting PDF file to DOCX file
        pdf.SaveToFile(f"pdf-docx/working/{pdf_name}/split_docxs/{split_pdf_path.split(".")[0]}.docx", FF.DOCX)
        pdf.Close()
    
    return f"pdf-docx/working/{pdf_name}/split_docxs" # return the path to the directory containing the split DOCXs

### step 3 - merging split DOCXs into one DOCX
def merge_docxs(split_docx_dir, pdf_name):
    os.makedirs(f"pdf-docx/working/{pdf_name}/temp-output", exist_ok = True)
    doc = Document()
    doc.LoadFromFile(f"pdf-docx/working/{pdf_name}/split_docxs/Split-1.docx")
    
    # inserting the content from other Word documents to first one
    for i in range(2, len(os.listdir(split_docx_dir)) + 1):
        doc.InsertTextFromFile(f"{split_docx_dir}/Split-{i}.docx", FileFormat.Auto)
    
    # saving the document
    doc.SaveToFile(f"pdf-docx/working/{pdf_name}/temp-output/{pdf_name}-spire.docx")
    doc.Close()
    return f"pdf-docx/working/{pdf_name}/temp-output/{pdf_name}-spire.docx" # return the path to the merged DOCX with spire watermarks

### step 4 - removing spire watermarks
def remove_watermarks(spire_docx_path, pdf_name):
    # defining regex patterns for the two evaluation warnings
    eval1 = re.compile("Evaluation Warning : The document was created with Spire.PDF for Python.")
    eval2 = re.compile("Evaluation Warning: The document was created with Spire.Doc for Python.")

    def paragraph_replace_text(paragraph, regex, replace_str):
        """
        Return `paragraph` after replacing all matches for `regex` with `replace_str`.
        """
        # a paragraph may contain more than one match, loop until all are replaced
        while True:
            text = paragraph.text
            match = regex.search(text)
            if not match:
                break

            # when there's a match, we need to modify run.text for each run that contains any part of the match-string.
            runs = iter(paragraph.runs)
            start, end = match.start(), match.end()

            # skip over any leading runs that do not contain the match
            for run in runs:
                run_len = len(run.text)
                if start < run_len:
                    break
                start, end = start - run_len, end - run_len

            # match starts somewhere in the current run. replace match-str prefix occurring in this run with entire replacement str.
            run_text = run.text
            run_len = len(run_text)
            run.text = "%s%s%s" % (run_text[:start], replace_str, run_text[end:])
            end -= run_len  # note this is run-len before replacement

            # remove any suffix of match word that occurs in following runs. note that such a suffix will always begin at the first character of the run. also note a suffix can span one or more entire following runs.
            for run in runs:  # next and remaining runs, uses same iterator
                if end <= 0:
                    break
                run_text = run.text
                run_len = len(run_text)
                run.text = run_text[end:]
                end -= run_len

        return paragraph

    document = docx_document(spire_docx_path)
    for paragraph in document.paragraphs:
        paragraph_replace_text(paragraph, eval1, "")
        paragraph_replace_text(paragraph, eval2, "")

    # saving the document with the evaluation warnings removed
    document.save(f"pdf-docx/working/{pdf_name}/{pdf_name}.docx")

### main
workingDir = "pdf-docx/working"
os.makedirs(workingDir, exist_ok = True)

# path to the directory containing PDF files to be converted
path = "pdf-docx/to-convert"
for pdf in os.listdir(path):
    pdf_name = pdf.split(".")[0]
    pdf_path = f"{path}/{pdf}"
    print(f"Splitting ({pdf_name})...")
    split_pdf_dir = split_pdf(pdf_path, pdf_name)
    print(f"({pdf_name}) split into multiple PDFs. Converting to DOCX...")
    split_docx_dir = convert_pdf(split_pdf_dir, pdf_name)
    print(f"({pdf_name})'s splits converted into DOCX. Merging into one DOCX...")
    spire_docx_path = merge_docxs(split_docx_dir, pdf_name)
    print(f"({spire_docx_path.split(".")[0]}) created. Removing watermarks...")
    remove_watermarks(spire_docx_path, pdf_name)
    print(f"({pdf_name}) converted to DOCX!\n")
