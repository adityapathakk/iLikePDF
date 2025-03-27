### importing libraries
import os
import re
from flask import Flask, request, send_file, jsonify
from io import BytesIO
from docx import Document as docx_document
from spire.pdf import PdfDocument
from spire.pdf import FileFormat as FF
from spire.doc import *
from spire.doc.common import *

# initializing Flask app
app = Flask(__name__)

# directory for storing PDF and DOCX files
working_dir = "pdf-docx-api/working"
os.makedirs(working_dir, exist_ok = True)

# step 1 - splitting PDFs
def split_pdf(pdf_path, pdf_name):
    os.makedirs(f"{working_dir}/{pdf_name}/split_pdfs", exist_ok=True)
    target_doc = PdfDocument()
    target_doc.LoadFromFile(pdf_path)
    
    num_required_splits = target_doc.Pages.Count // 10 + 1
    docs = [PdfDocument() for _ in range(num_required_splits)]
    
    doc_index = 0
    remainingPages = target_doc.Pages.Count
    current_page = 0
    while remainingPages >= 1:
        if remainingPages >= 10:
            docs[doc_index].InsertPageRange(target_doc, current_page, current_page + 9)
            current_page += 10
            remainingPages -= 10
            doc_index += 1
            continue
        elif remainingPages >= 1:
            docs[doc_index].InsertPageRange(target_doc, current_page, current_page + remainingPages - 1)
            doc_index += 1
            remainingPages = 0
            current_page += remainingPages
            break
    
    for i, doc in enumerate(docs):
        doc.SaveToFile(f"{working_dir}/{pdf_name}/split_pdfs/Split-{i + 1}.pdf")
        doc.Close()

    target_doc.Close()
    return f"{working_dir}/{pdf_name}/split_pdfs" # returning directory path containing the split PDFs

# step 2 - converting split PDFs to DOCXs
def convert_pdf(split_pdf_dir, pdf_name):
    os.makedirs(f"{working_dir}/{pdf_name}/split_docxs", exist_ok = True)
    for split_pdf_path in os.listdir(split_pdf_dir):
        pdf = PdfDocument()
        pdf.LoadFromFile(os.path.join(split_pdf_dir, split_pdf_path))
        
        pdf.SaveToFile(f"{working_dir}/{pdf_name}/split_docxs/{split_pdf_path.split('.')[0]}.docx", FF.DOCX)
        pdf.Close()
    
    return f"{working_dir}/{pdf_name}/split_docxs" # returning directory path containing the split DOCXs

# step 3 - merging split DOCXs into one DOCX
def merge_docxs(split_docx_dir, pdf_name):
    os.makedirs(f"{working_dir}/{pdf_name}/temp-output", exist_ok = True)
    doc = Document()
    doc.LoadFromFile(f"{working_dir}/{pdf_name}/split_docxs/Split-1.docx")
    
    for i in range(2, len(os.listdir(split_docx_dir)) + 1):
        doc.InsertTextFromFile(f"{split_docx_dir}/Split-{i}.docx", FileFormat.Auto)
    
    doc.SaveToFile(f"{working_dir}/{pdf_name}/temp-output/{pdf_name}-spire.docx")
    doc.Close()
    return f"{working_dir}/{pdf_name}/temp-output/{pdf_name}-spire.docx" # returning the path to the merged DOCX file with Spire watermarks

# step 4 - removing Spire watermarks
def remove_watermarks(spire_docx_path, pdf_name):
    eval1 = re.compile("Evaluation Warning : The document was created with Spire.PDF for Python.")
    eval2 = re.compile("Evaluation Warning: The document was created with Spire.Doc for Python.")

    def paragraph_replace_text(paragraph, regex, replace_str):
        while True:
            text = paragraph.text
            match = regex.search(text)
            if not match:
                break
            runs = iter(paragraph.runs)
            start, end = match.start(), match.end()
            for run in runs:
                run_len = len(run.text)
                if start < run_len:
                    break
                start, end = start - run_len, end - run_len
            run_text = run.text
            run_len = len(run_text)
            run.text = "%s%s%s" % (run_text[:start], replace_str, run_text[end:])
            end -= run_len

            for run in runs:
                if end <= 0:
                    break
                run_text = run.text
                run_len = len(run_text)
                run.text = run_text[end:]
                end -= run_len
        return paragraph

    document = docx_document(spire_docx_path)
    for paragraph in document.paragraphs:
        paragraph_replace_text(paragraph, eval1, "")
        paragraph_replace_text(paragraph, eval2, "")

    document.save(f"{working_dir}/{pdf_name}/{pdf_name}.docx") # saving the final DOCX file

# Flask route to handle PDF conversion
@app.route('/convert-pdf', methods = ['POST'])
def convert_pdf_api():
    try:
        # get the uploaded PDF file from request
        file = request.files['file']
        
        # save the uploaded file
        pdf_name = file.filename.split('.')[0]
        pdf_path = f"{working_dir}/{pdf_name}.pdf"
        file.save(pdf_path)
        
        # process the PDF file
        split_pdf_dir = split_pdf(pdf_path, pdf_name)
        split_docx_dir = convert_pdf(split_pdf_dir, pdf_name)
        spire_docx_path = merge_docxs(split_docx_dir, pdf_name)
        remove_watermarks(spire_docx_path, pdf_name)
        
        # return the final converted DOCX file
        docx_path = f"{working_dir}/{pdf_name}/{pdf_name}.docx"
        return send_file(docx_path, as_attachment = True, download_name = f"{pdf_name}.docx")

    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    app.run(debug = True)
